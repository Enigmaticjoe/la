#!/usr/bin/env python3
"""
Chimera RAG Processor - Knowledge Ingestion & Vector Database Manager
Processes documents from Unraid storage and creates searchable embeddings
"""

import os
import asyncio
import time
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime

from fastapi import FastAPI, BackgroundTasks, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from loguru import logger
import httpx
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct
from sqlalchemy import create_engine, Column, Integer, String, DateTime, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import redis

# Document processors
from pypdf import PdfReader
from docx import Document
import json

# Configure logging
logger.add("/logs/rag_processor.log", rotation="100 MB", retention="30 days", level="INFO")

# Configuration
OLLAMA_API = os.getenv("OLLAMA_API", "http://chimera_brain:11434")
QDRANT_URL = os.getenv("QDRANT_URL", "http://chimera_memory:6333")
POSTGRES_URL = os.getenv("POSTGRES_URL", "postgresql://chimera:chimera_secure_password_change_me@chimera_postgres:5432/chimera")
REDIS_URL = os.getenv("REDIS_URL", "redis://chimera_redis:6379")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "nomic-embed-text")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))
BATCH_SIZE = int(os.getenv("BATCH_SIZE", "32"))

DOCUMENTS_PATH = Path("/data/documents")
KNOWLEDGE_PATH = Path("/data/knowledge")

# FastAPI app
app = FastAPI(title="Chimera RAG Processor", version="1.0.0")

# Database setup
Base = declarative_base()

class ProcessedDocument(Base):
    __tablename__ = "processed_documents"

    id = Column(Integer, primary_key=True)
    filename = Column(String, unique=True, index=True)
    filepath = Column(String)
    filetype = Column(String)
    filesize = Column(Integer)
    chunk_count = Column(Integer)
    processed_at = Column(DateTime, default=datetime.utcnow)
    status = Column(String, default="pending")
    error_message = Column(Text, nullable=True)
    metadata_json = Column(Text)

engine = create_engine(POSTGRES_URL)
Base.metadata.create_all(engine)
SessionLocal = sessionmaker(bind=engine)

# Redis client for caching
redis_client = redis.from_url(REDIS_URL, decode_responses=True)

# Qdrant client
qdrant_client = QdrantClient(url=QDRANT_URL)

# Ensure collection exists
COLLECTION_NAME = "chimera_knowledge"
try:
    qdrant_client.get_collection(COLLECTION_NAME)
    logger.info(f"Collection '{COLLECTION_NAME}' already exists")
except:
    logger.info(f"Creating collection '{COLLECTION_NAME}'")
    qdrant_client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=768, distance=Distance.COSINE)
    )

# Models
class DocumentRequest(BaseModel):
    url: Optional[str] = None
    text: Optional[str] = None
    metadata: Dict = Field(default_factory=dict)

class SearchRequest(BaseModel):
    query: str
    limit: int = Field(default=10, ge=1, le=100)
    filter: Optional[Dict] = None

class HealthResponse(BaseModel):
    status: str
    timestamp: str
    services: Dict

# Document processors
def extract_text_from_pdf(filepath: Path) -> str:
    """Extract text from PDF file"""
    try:
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF: {e}")
        return ""

def extract_text_from_docx(filepath: Path) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        return ""

def extract_text_from_txt(filepath: Path) -> str:
    """Extract text from TXT file"""
    try:
        return filepath.read_text(encoding='utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Error reading TXT: {e}")
        return ""

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks"""
    if not text:
        return []

    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]

        # Try to end at sentence boundary
        if end < text_length:
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            boundary = max(last_period, last_newline)
            if boundary > chunk_size // 2:
                chunk = text[start:start + boundary + 1]
                end = start + boundary + 1

        chunks.append(chunk.strip())
        start = end - overlap

    return chunks

async def get_embedding(text: str) -> List[float]:
    """Get embedding vector from Ollama"""
    cache_key = f"embed:{hash(text)}"

    # Check cache
    cached = redis_client.get(cache_key)
    if cached:
        return json.loads(cached)

    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                f"{OLLAMA_API}/api/embeddings",
                json={"model": EMBEDDING_MODEL, "prompt": text},
                timeout=30.0
            )
            response.raise_for_status()
            embedding = response.json()["embedding"]

            # Cache for 24 hours
            redis_client.setex(cache_key, 86400, json.dumps(embedding))

            return embedding
        except Exception as e:
            logger.error(f"Error getting embedding: {e}")
            raise

async def process_document(filepath: Path) -> Dict:
    """Process a document and add to vector database"""
    logger.info(f"Processing document: {filepath}")

    db = SessionLocal()
    try:
        # Check if already processed
        existing = db.query(ProcessedDocument).filter_by(filename=filepath.name).first()
        if existing and existing.status == "completed":
            logger.info(f"Document already processed: {filepath.name}")
            return {"status": "skipped", "reason": "already_processed"}

        # Extract text based on file type
        suffix = filepath.suffix.lower()
        if suffix == '.pdf':
            text = extract_text_from_pdf(filepath)
        elif suffix in ['.docx', '.doc']:
            text = extract_text_from_docx(filepath)
        elif suffix in ['.txt', '.md', '.json']:
            text = extract_text_from_txt(filepath)
        else:
            logger.warning(f"Unsupported file type: {suffix}")
            return {"status": "error", "reason": "unsupported_type"}

        if not text or len(text) < 50:
            logger.warning(f"Insufficient text extracted from {filepath.name}")
            return {"status": "error", "reason": "insufficient_text"}

        # Chunk the text
        chunks = chunk_text(text)
        logger.info(f"Created {len(chunks)} chunks from {filepath.name}")

        # Generate embeddings and store
        points = []
        for idx, chunk in enumerate(chunks):
            embedding = await get_embedding(chunk)

            point = PointStruct(
                id=hash(f"{filepath.name}_{idx}"),
                vector=embedding,
                payload={
                    "text": chunk,
                    "filename": filepath.name,
                    "filepath": str(filepath),
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                    "filetype": suffix,
                    "processed_at": datetime.utcnow().isoformat(),
                }
            )
            points.append(point)

        # Batch insert to Qdrant
        qdrant_client.upsert(collection_name=COLLECTION_NAME, points=points)

        # Record in database
        doc_record = ProcessedDocument(
            filename=filepath.name,
            filepath=str(filepath),
            filetype=suffix,
            filesize=filepath.stat().st_size,
            chunk_count=len(chunks),
            status="completed",
            metadata_json=json.dumps({"chunks": len(chunks)})
        )

        if existing:
            db.delete(existing)

        db.add(doc_record)
        db.commit()

        logger.info(f"Successfully processed {filepath.name}")
        return {
            "status": "success",
            "filename": filepath.name,
            "chunks": len(chunks),
            "filesize": filepath.stat().st_size
        }

    except Exception as e:
        logger.error(f"Error processing {filepath}: {e}")
        db.rollback()
        return {"status": "error", "reason": str(e)}
    finally:
        db.close()

async def scan_and_process_documents():
    """Scan documents directory and process new files"""
    logger.info("Scanning documents directory...")

    if not DOCUMENTS_PATH.exists():
        logger.warning(f"Documents path does not exist: {DOCUMENTS_PATH}")
        return

    results = []
    for filepath in DOCUMENTS_PATH.rglob("*"):
        if filepath.is_file() and filepath.suffix.lower() in ['.pdf', '.docx', '.doc', '.txt', '.md', '.json']:
            result = await process_document(filepath)
            results.append(result)
            # Small delay to avoid overwhelming the system
            await asyncio.sleep(0.1)

    logger.info(f"Scan complete. Processed {len(results)} documents")
    return results

# API Endpoints
@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint"""
    services = {
        "ollama": "unknown",
        "qdrant": "unknown",
        "postgres": "unknown",
        "redis": "unknown"
    }

    # Check Ollama
    try:
        async with httpx.AsyncClient() as client:
            resp = await client.get(f"{OLLAMA_API}/api/tags", timeout=5.0)
            services["ollama"] = "healthy" if resp.status_code == 200 else "unhealthy"
    except:
        services["ollama"] = "unhealthy"

    # Check Qdrant
    try:
        qdrant_client.get_collections()
        services["qdrant"] = "healthy"
    except:
        services["qdrant"] = "unhealthy"

    # Check PostgreSQL
    try:
        db = SessionLocal()
        db.execute("SELECT 1")
        db.close()
        services["postgres"] = "healthy"
    except:
        services["postgres"] = "unhealthy"

    # Check Redis
    try:
        redis_client.ping()
        services["redis"] = "healthy"
    except:
        services["redis"] = "unhealthy"

    overall_status = "healthy" if all(s == "healthy" for s in services.values()) else "degraded"

    return HealthResponse(
        status=overall_status,
        timestamp=datetime.utcnow().isoformat(),
        services=services
    )

@app.post("/scan")
async def trigger_scan(background_tasks: BackgroundTasks):
    """Trigger document scanning"""
    background_tasks.add_task(scan_and_process_documents)
    return {"message": "Document scan started in background"}

@app.post("/search")
async def search_knowledge(request: SearchRequest):
    """Search the knowledge base"""
    try:
        # Get query embedding
        query_embedding = await get_embedding(request.query)

        # Search Qdrant
        search_results = qdrant_client.search(
            collection_name=COLLECTION_NAME,
            query_vector=query_embedding,
            limit=request.limit,
            query_filter=request.filter
        )

        # Format results
        results = []
        for hit in search_results:
            results.append({
                "text": hit.payload["text"],
                "filename": hit.payload["filename"],
                "chunk_index": hit.payload["chunk_index"],
                "score": hit.score,
                "metadata": {k: v for k, v in hit.payload.items() if k not in ["text"]}
            })

        return {"results": results, "count": len(results)}

    except Exception as e:
        logger.error(f"Search error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/stats")
async def get_stats():
    """Get processing statistics"""
    db = SessionLocal()
    try:
        total_docs = db.query(ProcessedDocument).count()
        completed = db.query(ProcessedDocument).filter_by(status="completed").count()
        failed = db.query(ProcessedDocument).filter_by(status="error").count()

        collection_info = qdrant_client.get_collection(COLLECTION_NAME)

        return {
            "documents": {
                "total": total_docs,
                "completed": completed,
                "failed": failed
            },
            "vectors": {
                "count": collection_info.points_count,
                "collection": COLLECTION_NAME
            }
        }
    finally:
        db.close()

@app.on_event("startup")
async def startup_event():
    """Run initial scan on startup"""
    logger.info("RAG Processor starting up...")
    logger.info(f"Documents path: {DOCUMENTS_PATH}")
    logger.info(f"Knowledge path: {KNOWLEDGE_PATH}")
    logger.info(f"Ollama API: {OLLAMA_API}")
    logger.info(f"Qdrant URL: {QDRANT_URL}")

    # Wait for services
    await asyncio.sleep(10)

    # Trigger initial scan
    asyncio.create_task(scan_and_process_documents())

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
